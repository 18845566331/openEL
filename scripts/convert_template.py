from docx import Document
import re
import os

def convert_template(input_path, output_path):
    doc = Document(input_path)
    pattern = re.compile(r"\{([^{}]+)\}")
    
    def process_text(text):
        # Replace {key} with {{ key }}
        # But be careful not to double replace if already {{ }}
        # The user template has {key}.
        return pattern.sub(r"{{\1}}", text)

    for p in doc.paragraphs:
        if pattern.search(p.text):
            # Replacing in text might lose formatting if runs are split
            # But for simple placeholders it often works.
            # A better approach for docx is iterating runs, but simpler regex usually suffices for placeholders
            p.text = process_text(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if pattern.search(p.text):
                        p.text = process_text(p.text)
                        
    doc.save(output_path)
    print(f"Converted template saved to {output_path}")

if __name__ == "__main__":
    input_file = r"D:\opencv缺陷检测 - 副本\NIT NNE RT-002 光伏组件EL检测报告模板.docx"
    output_temp = r"D:\opencv缺陷检测 - 副本\el_defect_system\backend\app\template_temp.docx"
    convert_template(input_file, output_temp)
