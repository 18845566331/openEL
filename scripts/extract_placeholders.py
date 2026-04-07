from docx import Document
import re
import os

template_path = r"D:\opencv缺陷检测 - 副本\NIT NNE RT-002 光伏组件EL检测报告模板.docx"

if not os.path.exists(template_path):
    print(f"Error: Template not found at {template_path}")
    exit(1)

doc = Document(template_path)
placeholders = set()
pattern = re.compile(r"\{([^{}]+)\}")

def check_text(text):
    matches = pattern.findall(text)
    for m in matches:
        placeholders.add(m)

# Check paragraphs
for p in doc.paragraphs:
    check_text(p.text)

# Check tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                check_text(p.text)

print("Found placeholders:")
for p in sorted(list(placeholders)):
    print(f"- {p}")
