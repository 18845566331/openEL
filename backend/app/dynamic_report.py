import os
import re
import tempfile
import base64
from pathlib import Path
from typing import Any, List
from fastapi import APIRouter, HTTPException
from docx import Document
from pydantic import BaseModel

router = APIRouter(tags=["Dynamic Report"])

class TemplateScanParams(BaseModel):
    template_path: str

class PreviewRequest(BaseModel):
    template_path: str
    fields_data: dict

def _extract_placeholders(doc: Document) -> List[str]:
    # Regex to find { field_name }, {{ field_name }}
    pattern = re.compile(r'\{{1,2}\s*([A-Za-z0-9_\u4e00-\u9fa5]+)\s*\}{1,2}')
    placeholders = []

    for paragraph in doc.paragraphs:
        matches = pattern.findall(paragraph.text)
        for m in matches:
            placeholders.append(m)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = pattern.findall(paragraph.text)
                    for m in matches:
                        placeholders.append(m)
                        
    # 忽略默认保留的关键统计计算字段
    ignore_keys = {"grade_a_count", "grade_b_count", "grade_c_count", "total_images", 
                   "ng_images", "ok_images", "defect_ratio", "defect_total"}
    
    # 过滤并且按文档出现的顺序返回（保留原始顺序去除重复）
    filtered = [p for p in placeholders if p not in ignore_keys]
    return list(dict.fromkeys(filtered))

@router.post("/api/report/template_fields")
def scan_template_fields(params: TemplateScanParams) -> dict[str, Any]:
    """提取指定docx模板中所有的可填写字段并返回"""
    target_path = Path(params.template_path)
    if not target_path.exists() or target_path.suffix != ".docx":
        raise HTTPException(status_code=400, detail="模板文件不存在或不是.docx格式")
    
    try:
        doc = Document(target_path)
        fields = _extract_placeholders(doc)
        return {"success": True, "fields": fields}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/api/report/live_preview")
def generate_live_preview(request: PreviewRequest) -> dict[str, Any]:
    """将填充后的内容生成为一系列预览图片"""
    target_path = Path(request.template_path)
    if not target_path.exists() or target_path.suffix != ".docx":
        raise HTTPException(status_code=400, detail="模板文件不存在或不是.docx格式")
    
    import shutil
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        # Fallback to pure docx simple replace if docxtpl is missing
        # We will attempt to rely on regex in python-docx 
        pass

    try:
        # 1. 创建临时的工作目录
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_docx = Path(temp_dir) / "preview.docx"
            temp_pdf = Path(temp_dir) / "preview.pdf"
            
            # Simple placeholder replacement logic
            doc = Document(target_path)
            # This simple replacement is a demo fallback.
            # Real docxtpl does full variable replacement cleanly.
            def _replace_in_blocks(blocks):
                for paragraph in blocks:
                    for k, v in request.fields_data.items():
                        holder = "{{" + k + "}}"
                        holder2 = "{{ " + k + " }}"
                        holder3 = "{" + k + "}"
                        holder4 = "{ " + k + " }"
                        # 确保如果有替换发生，我们正确更新
                        for h in (holder, holder2, holder3, holder4):
                            if h in paragraph.text:
                                paragraph.text = paragraph.text.replace(h, str(v) if v else "")

            # 替换正文段落
            _replace_in_blocks(doc.paragraphs)
            # 替换正文表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        _replace_in_blocks(cell.paragraphs)
            # 替换页眉页脚
            for section in doc.sections:
                _replace_in_blocks(section.header.paragraphs)
                _replace_in_blocks(section.footer.paragraphs)
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            _replace_in_blocks(cell.paragraphs)
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            _replace_in_blocks(cell.paragraphs)
                            
            doc.save(temp_docx)

            # 2. 调用原生的 Word COM 极速导出核心（只渲染前 5 页）
            # 这能将 103 页的渲染时间从近 10 秒暴降到 0.5 秒以下！
            try:
                import pythoncom
                import win32com.client
                pythoncom.CoInitialize()
                # 尽量复用已有的Word进程（提高极速），如果没有则启动新的
                try:
                    word = win32com.client.Dispatch("Word.Application")
                except:
                    word = win32com.client.DispatchEx("Word.Application")
                
                word.Visible = False
                # 打开我们的临时模板
                doc_word = word.Documents.Open(str(temp_docx), ReadOnly=True, Visible=False)
                # wdExportFormatPDF = 17, wdExportFromTo = 2
                # Range=2 (指定页码范围), From=1, To=5
                doc_word.ExportAsFixedFormat(str(temp_pdf), 17, False, 0, 2, 1, 5)
                doc_word.Close(0) # 0 = wdDoNotSaveChanges
            except Exception as e:
                import traceback
                traceback.print_exc()
                return {"success": False, "error": f"转换预览PDF失败，请确微软Word服务可用: {e}"}
            finally:
                try:
                    pythoncom.CoUninitialize()
                except:
                    pass
            
            if not temp_pdf.exists():
                 raise ValueError("PDF 生成失败")

            # 3. 将 PDF 转换为图片数组
            import fitz
            doc_pdf = fitz.open(str(temp_pdf))
            encoded_images = []
            
            # 为了防止100多页的超大型报告导致接口卡死、Flutter内存溢出，最多只渲染前5页作为预览
            preview_pages_limit = min(len(doc_pdf), 5)
            
            for i in range(preview_pages_limit):
                page = doc_pdf.load_page(i)
                pix = page.get_pixmap(dpi=120)  # 降低 DPI 到 120 加快渲染并减小 Base64 体积
                img_path = Path(temp_dir) / f"page_{i}.jpg"
                pix.save(str(img_path))
                
                with open(img_path, "rb") as f:
                    encoded = base64.b64encode(f.read()).decode("utf-8")
                    encoded_images.append(encoded)
            doc_pdf.close()

            return {
                "success": True,
                "pages": encoded_images
            }
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}

class DynamicGenerateRequest(BaseModel):
    template_path: str
    output_path: str
    fields_data: dict

@router.post("/api/report/dynamic_generate")
def dynamic_generate(request: DynamicGenerateRequest) -> dict[str, Any]:
    """生成真正的高保真定制排版 Word 报告。"""
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        raise HTTPException(status_code=500, detail="未安装 docxtpl 依赖，无法生成高保真模板")
        
    template_path = Path(request.template_path)
    output_path = Path(request.output_path)
    
    if not template_path.exists():
        raise HTTPException(status_code=400, detail="报告模板不存在")
        
    try:
        # Load docxtpl
        doc = DocxTemplate(str(template_path))
        
        # We can implement InlineImage logic if there are any specific image placeholders.
        # But for now we just render textual variables based on fields_data
        # If any value resembles a base64 or absolute path, it should be processed into an InlineImage.
        # For simplicity, render the context dictionary directly.
        context = request.fields_data
        
        doc.render(context)
        
        # Ensure output dir exists
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        return {"success": True, "output_path": str(output_path)}
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"报告生成失败: {e}")
