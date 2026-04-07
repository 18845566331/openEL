import os
import io
import time
import json
import uuid
import shutil
import base64
import zipfile
import urllib.parse
from typing import Any, List, Optional
from fastapi import APIRouter, HTTPException, File, UploadFile, Form, BackgroundTasks
from fastapi.responses import JSONResponse, Response, FileResponse, StreamingResponse
from pydantic import BaseModel
import cv2
import numpy as np
import requests as _requests_lib
import logging
from datetime import datetime

logger = logging.getLogger(__name__)



router = APIRouter()

from pathlib import Path
from app.schemas import *
from app.detector import DetectionItem
from app.state import _loaded_model_id, _model_profiles_cache, _cache_lock, engine


import docx
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import openpyxl
from openpyxl.drawing.image import Image as ExcelImage
from PIL import Image
from app.main import get_batch_status 
def export_csv_report(request: dict[str, Any]) -> dict[str, Any]:
    """导出 CSV 检测报告。"""
    project_info = request.get("project_info", {})
    output_path = request.get("output_path", "")
    if not output_path:
        raise HTTPException(status_code=400, detail="未指定输出路径")
    logger.info("CSV报告导出开始: output=%s", output_path)
    try:
        target = Path(output_path).expanduser().resolve()
        target.parent.mkdir(parents=True, exist_ok=True)
        file_results = project_info.get("file_results", [])
        defect_by_class = project_info.get("defect_by_class", {})

        with target.open("w", newline="", encoding="utf-8-sig") as file:
            writer = csv.writer(file)

            def _g(*keys: str, default: str = "") -> str:
                for k in keys:
                    val = project_info.get(k)
                    if val is not None and str(val).strip():
                        return str(val).strip()
                return default

            # 项目信息头
            writer.writerow(["项目名称", _g("项目名称", "project_name")])
            writer.writerow(["报告编号", _g("报告编号", "report_code")])
            writer.writerow(["委托单位", _g("委托单位", "client_name")])
            writer.writerow(["检测单位", _g("检测单位", "test_unit", "testing_unit")])
            writer.writerow(["检测人员", _g("检测人员", "tester")])
            writer.writerow(["签发日期", _g("签发日期", "issue_date", "signing_date")])
            writer.writerow([])
            # 缺陷统计
            writer.writerow(["缺陷类别", "数量"])
            for cls_name, count in defect_by_class.items():
                writer.writerow([cls_name, count])
            writer.writerow([])
            writer.writerow(["文件名", "缺陷类型", "纬度(Lat)", "经度(Lon)", "文件路径"])
            for f in file_results:
                fname = f.get("name", "")
                fpath = f.get("path", "")
                detections = f.get("detections", [])
                if not detections:
                    writer.writerow([fname, "OK", "", "", fpath])
                else:
                    for det in detections:
                        cname = det.get("class_name", "")
                        lat = det.get("lat", "")
                        lon = det.get("lon", "")
                        writer.writerow([fname, cname, lat, lon, fpath])

        return {"message": "CSV 报告导出成功", "output_path": target.as_posix()}
    except Exception as exc:
        logger.error("CSV导出失败: %s", exc, exc_info=True)
        raise HTTPException(status_code=400, detail=f"CSV导出失败: {exc}") from exc

@router.post("/api/report/export_geojson")
def export_geojson_report(request: dict[str, Any]) -> dict[str, Any]:
    """导出 GeoJSON 格式检测结果（用于 GIS 地图分析）。"""
    project_info = request.get("project_info", {})
    output_path = request.get("output_path", "")
    if not output_path:
        raise HTTPException(status_code=400, detail="未指定输出路径")

    logger.info("GeoJSON 导出开始: output=%s", output_path)
    try:
        import json
        target = Path(output_path).expanduser().resolve()
        target.parent.mkdir(parents=True, exist_ok=True)
        file_results = project_info.get("file_results", [])
        
        features = []
        for f in file_results:
            fname = f.get("name", "")
            fpath = f.get("path", "")
            detections = f.get("detections", [])
            for det in detections:
                lat = det.get("lat")
                lon = det.get("lon")
                if lat is not None and lon is not None:
                    features.append({
                        "type": "Feature",
                        "geometry": {
                            "type": "Point",
                            "coordinates": [lon, lat]
                        },
                        "properties": {
                            "file_name": fname,
                            "file_path": fpath,
                            "defect_type": det.get("class_name", ""),
                            "confidence": det.get("score", 0.0)
                        }
                    })

        geojson = {
            "type": "FeatureCollection",
            "features": features
        }
        
        with target.open("w", encoding="utf-8") as file:
            json.dump(geojson, file, ensure_ascii=False, indent=2)
            
        return {"message": "GeoJSON 导出成功", "output_path": target.as_posix()}
    except Exception as exc:
        logger.error("GeoJSON 导出失败: %s", exc, exc_info=True)
        raise HTTPException(status_code=400, detail=f"GeoJSON导出失败: {exc}") from exc


@router.post("/api/report/export_word")
def export_word_report(request: dict[str, Any]) -> dict[str, Any]:
    """导出 Word 检测报告（基于 NIT NNE RT-002 模板）。"""
    project_info = request.get("project_info", {})
    output_path = request.get("output_path", "")
    if not output_path:
        raise HTTPException(status_code=400, detail="未指定输出路径")
    
    # 强制增加扩展名
    if not output_path.lower().endswith('.docx'):
        output_path += '.docx'
        
    logger.info("Word报告导出开始: output=%s", output_path)
    logger.info("Word报告数据字段(%d): %s", len(project_info), list(project_info.keys())[:20])
    # DEBUG: 记录file_results数量以便调试200张图片限制问题
    _fr_count = len(project_info.get("file_results", []))
    logger.info("DEBUG: project_info中file_results数量=%d", _fr_count)
    for k in ['项目名称', 'project_name', '报告编号', '委托单位']:
        logger.info("  %s = %s", k, str(project_info.get(k, '<空>'))[:60])

    # 使用模板导出 — 按优先级搜索模板文件
    template_path = Path("__nonexistent__")  # 占位
    _template_candidates = []

    # 1. PyInstaller 打包后的 templates 目录（_MEIPASS 或 exe 同级）
    _bundle_dir = getattr(sys, '_MEIPASS', None)
    if _bundle_dir:
        _template_candidates.append(Path(_bundle_dir) / "templates")
    # 2. backend.exe 所在目录的 templates 子目录（--onedir 模式）
    _exe_dir = Path(sys.executable).parent if getattr(sys, 'frozen', False) else None
    if _exe_dir:
        _template_candidates.append(_exe_dir / "templates")
        _template_candidates.append(_exe_dir)
    # 3. 源码开发时：项目根目录（backend 的上级）
    _template_candidates.append(Path(__file__).resolve().parent.parent.parent)
    _template_candidates.append(Path(__file__).resolve().parent.parent)
    # 4. 当前工作目录
    _template_candidates.append(Path.cwd())

    for _cand_dir in _template_candidates:
        if not _cand_dir.is_dir():
            continue
        _found = list(_cand_dir.rglob("*光伏组件EL检测报告模板*.docx"))
        if not _found:
            _found = list(_cand_dir.rglob("NIT NNE RT-002*.docx"))  # 兼容旧模板名
        if _found:
            template_path = _found[0]
            logger.info("找到报告模板: %s", template_path)
            break

    if template_path.exists():
        return _export_word_from_template(project_info, output_path, template_path)
    else:
        logger.warning("未找到报告模板，使用程序化生成")
        return _export_word_fallback(request, output_path)

def _replace_placeholder_in_runs(runs, placeholder: str, value: str):
    """跨 run 替换占位符。

    Word 经常把 {项目名称} 拆成多个 run，如 '{' + '项目名称' + '}'。
    本函数将所有 run 的文本拼接后做替换，再重新分配回 run。
    保留第一个 run 的格式，清空其余 run。
    """
    if not runs:
        return
    full_text = "".join(r.text for r in runs)
    if placeholder not in full_text:
        return
    new_text = full_text.replace(placeholder, value)
    # 把新文本放到第一个 run，清空其余
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def _replace_in_paragraph(paragraph, placeholder: str, value: str):
    """在段落中替换占位符（跨 run 安全）。"""
    _replace_placeholder_in_runs(paragraph.runs, placeholder, value)


def _replace_in_cell(cell, placeholder: str, value: str):
    """在单元格的所有段落中替换占位符。"""
    for p in cell.paragraphs:
        _replace_in_paragraph(p, placeholder, value)


def _compress_image_for_docx(image_path: str, max_dim: int = 3000, quality: int = 85) -> str | None:
    """Compress/resize an image for Word insertion. Returns temp file path or None on failure."""
    try:
        from PIL import Image as PILImage
        img = PILImage.open(image_path)
        w, h = img.size
        needs_resize = w > max_dim or h > max_dim
        # Also compress if file is large (>5MB)
        needs_compress = Path(image_path).stat().st_size > 5 * 1024 * 1024
        if not needs_resize and not needs_compress:
            return None  # No compression needed, caller uses original
        if needs_resize:
            ratio = min(max_dim / w, max_dim / h)
            img = img.resize((int(w * ratio), int(h * ratio)), PILImage.LANCZOS)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        tmp = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
        img.save(tmp.name, "JPEG", quality=quality)
        tmp.close()
        logger.info("图片压缩: %s -> %s (%dx%d)", image_path, tmp.name, img.size[0], img.size[1])
        return tmp.name
    except Exception as exc:
        logger.warning("图片压缩失败 (%s): %s, 使用原图", type(exc).__name__, exc)
        return None


def _safe_remove(path: str):
    """Silently remove a file."""
    try:
        os.remove(path)
    except OSError:
        pass


def _export_word_from_template(
    project_info: dict[str, Any],
    output_path: str,
    template_path: Path,
) -> dict[str, Any]:
    """基于 NIT NNE RT-002 模板导出 Word 报告。

    处理策略：
    1. 跨 run 合并替换所有 {占位符}
    2. 图片占位符替换为实际图片
    3. Table 7 (EL检测数据) 动态填充检测结果行
    4. Table 6 缺陷统计行填充
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.oxml.ns import qn
    import copy

    logger.info("模板导出开始: template=%s, output=%s", template_path, output_path)

    def _g(*keys: str, default: str = "") -> str:
        for k in keys:
            val = project_info.get(k)
            if val is not None and str(val).strip():
                return str(val).strip()
        return default

    doc = Document(str(template_path))

    # ── 1. 构建占位符 → 值 映射表 ──
    mapping = {
        "{报告编号}": _g("报告编号", "report_code"),
        "{签发日期}": _g("签发日期", "issue_date", "signing_date"),
        "{委托单位}": _g("委托单位", "client_name"),
        "{检测单位}": _g("检测单位", "test_unit", "testing_unit"),
        "{项目名称}": _g("项目名称", "project_name"),
        "{项目地址}": _g("项目地址", "project_address"),
        "{委托单位地址}": _g("委托单位地址", "client_address"),
        "{样品来源}": _g("样品来源", "sample_source"),
        "{抽样原则}": _g("抽样原则", "sampling_principle"),
        "{抽样原则 }": _g("抽样原则", "sampling_principle"),  # 模板中有尾部空格
        "{参考标准}": _g("参考标准", "reference_standard"),
        "{项目概述}": _g("项目概述", "project_overview"),
        "{ YYYY/MM/DD~ YYYY/MM/DD }": _g("尽调周期", "diligence_period", "investigation_period"),
        "{YYYY/MM/DD~ YYYY/MM/DD}": _g("尽调周期", "diligence_period", "investigation_period"),
        "{电站地址}": _g("电站地址", "station_address"),
        "{联系人}": _g("联系人", "contact_person"),
        "{联系方式}": _g("联系方式", "contact_phone"),
        "{业主单位名称}": _g("业主单位名称", "owner_name"),
        "{设计单位名称}": _g("设计单位名称", "design_unit"),
        "{EPC单位名称}": _g("EPC单位名称", "epc_unit"),
        "{运维单位名称}": _g("运维单位名称", "maintenance_unit"),
        "{电站直流安装容量}": _g("电站直流安装容量", "dc_capacity"),
        "{分几期建设}": _g("分几期建设"),
        "{状态}": _g("建设状态", "状态"),
        "{电站类型}": _g("电站类型"),
        "{土地类型}": _g("土地类型"),
        "{土地现状}": _g("土地现状"),
        "{占地面积}": _g("占地面积"),
        "{建设成本}": _g("建设成本"),
        "{地形地貌}": _g("地形地貌"),
        "{水/土壤情况}": _g("水/土壤情况"),
        "{电网接入方式}": _g("电网接入方式"),
        "{并网电压等级}": _g("并网电压等级"),
        "{并网接入距离}": _g("并网接入距离"),
        "{主变容量}": _g("主变容量"),
        "{是否限电}": _g("是否限电"),
        "{限电调控方式}": _g("限电调控方式"),
        "{上网电价}": _g("上网电价"),
        "{桩基形式}": _g("桩基形式"),
        "{支架形式}": _g("支架形式"),
        "{单个组串组件数}": _g("单个组串组件数"),
        "{组件固定方式}": _g("组件固定方式"),
        "{组件下边缘距地高度}": _g("组件下边缘距地高度"),
        "{是否配置气象站}": _g("是否配置气象站"),
        "{气象站距离光伏区距离}": _g("气象站距离光伏区距离"),
        "{气象站厂家}": _g("气象站厂家"),
        "{气象站采集数据类型}": _g("气象站采集数据类型"),
        # 光伏组件参数通过 section 4b 按位置填充，不使用占位符替换
        "{光伏厂区地理位置图}": "",  # 图片单独处理
        "{光伏厂区整体图}": "",  # 图片单独处理
        "{附件1 光伏组件电致发光（EL）检测数据}": "",  # EL表格插入点，单独处理
    }

    # ── 2. 替换段落中的占位符 ──
    for p in doc.paragraphs:
        full = "".join(r.text for r in p.runs)
        for ph, val in mapping.items():
            if ph in full:
                _replace_in_paragraph(p, ph, val)
                full = "".join(r.text for r in p.runs)

    # ── 3. 替换表格中的占位符 ──
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full = "".join(r.text for r in p.runs)
                    for ph, val in mapping.items():
                        if ph in full:
                            _replace_in_paragraph(p, ph, val)
                            full = "".join(r.text for r in p.runs)

    # ── 缓存原始表格引用（在 add_table 之前）──
    # doc.add_table() 会向 doc.tables 追加新表格，导致索引偏移。
    # 必须在任何 add_table 调用之前保存对原始表格的引用。
    _orig_tables = list(doc.tables)

    # ── 4. 处理图片占位符 (Table 3) ──
    try:
        if len(_orig_tables) > 3:
            img_table = _orig_tables[3]
            geo_img = _g("geographic_image_path")
            overall_img = _g("overall_image_path")
            # Row 0: 光伏厂区地理位置图
            if geo_img and Path(geo_img).exists():
                cell = img_table.rows[0].cells[0]
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run()
                compressed = _compress_image_for_docx(geo_img)
                run.add_picture(compressed or geo_img, width=Inches(5.5))
                if compressed:
                    _safe_remove(compressed)
            # Row 2: 光伏厂区整体图
            if overall_img and Path(overall_img).exists():
                cell = img_table.rows[2].cells[0]
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run()
                compressed = _compress_image_for_docx(overall_img)
                run.add_picture(compressed or overall_img, width=Inches(5.5))
                if compressed:
                    _safe_remove(compressed)
    except Exception as exc:
        logger.warning("图片插入失败 (%s): %s", type(exc).__name__, exc)

    # ── 4b. 填充 光伏组件参数（按标签匹配，直接写入值单元格）──
    # 遍历所有表格，找到包含"光伏组件"的表格，
    # 然后按标签单元格内容匹配字段名，将 pv_modules 数据写入对应的值单元格。
    try:
        pv_modules = project_info.get("pv_modules", [])
        logger.info("光伏组件数据: pv_modules数量=%d, 类型=%s",
                     len(pv_modules) if isinstance(pv_modules, list) else -1,
                     type(pv_modules).__name__)
        if pv_modules and isinstance(pv_modules, list):
            # 标签文本 → pv_module 字段名的映射
            # 标签单元格可能包含全角或半角括号，这里列出所有可能的标签文本
            _label_to_field = {}
            _field_names = [
                "生产厂家", "型号", "类型",
                "Pmax（Wp）", "Voc（V）", "Vmp（V）",
                "Isc（A）", "Imp（A）", "组件尺寸（mm）",
                "短路电流温度系数（%/°C）", "功率温度系数（%/°C）",
                "开路电压温度系数（%/°C）",
            ]
            for fn in _field_names:
                # 原始名
                _label_to_field[fn] = fn
                # 去掉所有空格的版本
                _label_to_field[fn.replace(" ", "")] = fn
                # 全角→半角
                hf = fn.replace("（", "(").replace("）", ")")
                _label_to_field[hf] = fn
                _label_to_field[hf.replace(" ", "")] = fn
                # 全角→半角+空格
                hfs = fn.replace("（", " (").replace("）", ")")
                _label_to_field[hfs] = fn
                _label_to_field[hfs.replace(" ", "")] = fn

            # 查找包含 "光伏组件" 的表格
            pv_table = None
            for t in _orig_tables:
                all_text = "".join("".join(c.text for c in row.cells) for row in t.rows)
                if "光伏组件" in all_text and "生产厂家" in all_text:
                    pv_table = t
                    break
            
            if not pv_table:
                # 备选：找包含"光伏组件"的表格
                for t in _orig_tables:
                    all_text = "".join("".join(c.text for c in row.cells) for row in t.rows)
                    if "光伏组件" in all_text:
                        pv_table = t
                        break

            if pv_table:
                # 找到每个 "光伏组件 N" 标题行的索引
                device_sections = []  # [(title_row_idx, device_index)]
                for ri, row in enumerate(pv_table.rows):
                    for c in row.cells:
                        ct = c.text.strip().replace(" ", "").replace("\u3000", "")
                        if ct.startswith("光伏组件") and len(ct) <= 10:
                            device_sections.append(ri)
                            break

                logger.info("光伏组件参数表: 总行数=%d, 设备标题行=%s, pv_modules数=%d",
                            len(pv_table.rows), device_sections, len(pv_modules))

                # 遍历每个设备区域，按标签匹配填充
                for dev_idx, module in enumerate(pv_modules):
                    if not isinstance(module, dict):
                        continue
                    if dev_idx >= len(device_sections):
                        logger.warning("模板中没有光伏组件 %d 的区域，跳过", dev_idx + 1)
                        continue
                    
                    title_row_idx = device_sections[dev_idx]
                    # 确定数据行范围：从标题行+1到下一个标题行（或表格末尾）
                    next_title = device_sections[dev_idx + 1] if dev_idx + 1 < len(device_sections) else len(pv_table.rows)
                    
                    filled_count = 0
                    for ri in range(title_row_idx + 1, next_title):
                        row = pv_table.rows[ri]
                        if len(row.cells) < 4:
                            continue
                        # 处理左侧：cells[0]=标签, cells[1]=值
                        # 处理右侧：cells[2]=标签, cells[3]=值
                        for label_ci, value_ci in [(0, 1), (2, 3)]:
                            label_text = row.cells[label_ci].text.strip().replace(" ", "").replace("\u3000", "")
                            if label_text in _label_to_field:
                                field_name = _label_to_field[label_text]
                                val = _pv_get_module_value(module, field_name)
                                if val:
                                    # 强制清除单元格所有内容，写入新值并设置格式
                                    cell = row.cells[value_ci]
                                    for p in cell.paragraphs:
                                        for r in p.runs:
                                            r.text = ""
                                    p0 = cell.paragraphs[0] if cell.paragraphs else None
                                    if p0 is not None:
                                        if p0.runs:
                                            p0.runs[0].text = val
                                            run = p0.runs[0]
                                        else:
                                            run = p0.add_run(val)
                                        # 字体：Times New Roman，五号（10.5pt）
                                        run.font.name = 'Times New Roman'
                                        run.font.size = Pt(10.5)
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                                        # 段落：居中对齐
                                        p0.alignment = 1  # CENTER
                                        # 段前段后 0.5 行（单倍行距下约 6pt）
                                        pf = p0.paragraph_format
                                        pf.space_before = Pt(6)
                                        pf.space_after = Pt(6)
                                        # 单倍行距
                                        pf.line_spacing = 1.0
                                    filled_count += 1
                                else:
                                    # 没有值，清除占位符花括号
                                    cell = row.cells[value_ci]
                                    cell_text = "".join(r.text for p in cell.paragraphs for r in p.runs)
                                    if "{" in cell_text:
                                        for p in cell.paragraphs:
                                            for r in p.runs:
                                                r.text = ""
                    
                    logger.info("光伏组件 %d: 填充了 %d 个字段", dev_idx + 1, filled_count)
            else:
                logger.warning("未找到光伏组件参数表")
    except Exception as exc:
        logger.warning("光伏组件参数填充失败 (%s): %s", type(exc).__name__, exc, exc_info=True)

    # ── 5. 填充缺陷统计表 ──
    # 通过内容搜索定位表格，用占位符关键字匹配填充（不依赖硬编码行列索引）
    try:
        t6 = None
        for t in _orig_tables:
            table_text = "".join(c.text for row in t.rows for c in row.cells).replace(" ", "").replace("　", "")
            if "缺陷名称" in table_text and "A类" in table_text and "C类" in table_text:
                t6 = t
                break

        if t6:
            defect_by_class = project_info.get("defect_by_class", {})
            defect_grade_breakdown = project_info.get("defect_grade_breakdown", {})
            total_images = project_info.get("total_images", 0)
            total_tested = max(1, total_images)

            # 优先使用 grade_a/b/c_count，但如果全为0则回退到 ok/ng_images
            _ga = project_info.get("grade_a_count", 0) or 0
            _gb = project_info.get("grade_b_count", 0) or 0
            _gc = project_info.get("grade_c_count", 0) or 0
            if _ga == 0 and _gb == 0 and _gc == 0:
                grade_a = project_info.get("ok_images", 0) or 0
                grade_b = 0
                grade_c = project_info.get("ng_images", 0) or 0
            else:
                grade_a = _ga
                grade_b = _gb
                grade_c = _gc

            pct_a = f"{grade_a / total_tested * 100:.1f}" if total_tested else "0"
            pct_b = f"{grade_b / total_tested * 100:.1f}" if total_tested else "0"
            pct_c = f"{grade_c / total_tested * 100:.1f}" if total_tested else "0"

            # ── 5a. 逐单元格扫描，通过占位符关键字填充 ──
            for row in t6.rows:
                _seen_tc = set()
                for cell in row.cells:
                    _tc_id = id(cell._tc)
                    if _tc_id in _seen_tc:
                        continue
                    _seen_tc.add(_tc_id)
                    ct = cell.text.replace(" ", "").replace("\u3000", "").replace("\n", "")

                    # 组件型号
                    if "{组件型号}" in ct:
                        model_val = _g("组件型号", "module_model", "型号")
                        if not model_val:
                            _pvm = project_info.get("pv_modules", [])
                            if isinstance(_pvm, list) and _pvm:
                                _models = [str(m.get("型号", "") or m.get("model", "")).strip()
                                           for m in _pvm if isinstance(m, dict)]
                                model_val = "、".join(m for m in _models if m)
                        _set_cell_text(cell, model_val)

                    # 被测组件数量
                    elif "{被测组件数量}" in ct:
                        _set_cell_text(cell, str(total_images))

                    # A类组件数量
                    elif "{A类组件数量" in ct and "比例" not in ct:
                        _set_cell_text(cell, str(grade_a))

                    # B类组件数量
                    elif "{B类组件数量" in ct and "比例" not in ct:
                        _set_cell_text(cell, str(grade_b))

                    # C类组件数量
                    elif "{C类组件数量" in ct and "比例" not in ct:
                        _set_cell_text(cell, str(grade_c))

                    # A类组件比例
                    elif "{A类组件比例" in ct:
                        _set_cell_text(cell, pct_a)

                    # B类组件比例
                    elif "{B类组件比例" in ct:
                        _set_cell_text(cell, pct_b)

                    # C类组件比例
                    elif "{C类组件比例" in ct:
                        _set_cell_text(cell, pct_c)

                    # 检测日期
                    elif "YYYY/MM/DD" in ct and "{" in ct:
                        _date_val = _g("尽调周期", "diligence_period", "检测日期", "detection_date")
                        _set_cell_text(cell, _date_val)

                    # 检测结果汇总
                    elif "{检测结果汇总}" in ct:
                        # 计算无缺陷图片数量
                        _fr = project_info.get("file_results", [])
                        ok_no_defect = sum(1 for f in _fr if isinstance(f, dict) and not f.get("detections"))
                        ok_pct = f"{ok_no_defect / total_tested * 100:.1f}" if total_tested else "0"
                        summary = (f"共检测组件{total_images}块，"
                                   f"A类{grade_a}块({pct_a}%)、"
                                   f"B类{grade_b}块({pct_b}%)、"
                                   f"C类{grade_c}块({pct_c}%)，"
                                   f"无缺陷组件{ok_no_defect}块({ok_pct}%)")
                        _set_cell_text(cell, summary)

            logger.info("缺陷统计汇总行填充完成: total=%d, A=%d, B=%d, C=%d", total_images, grade_a, grade_b, grade_c)

            # ── 5b. 从 file_results 计算每种缺陷的 A/B/C 等级统计 ──
            # 每张光伏板只归属一个缺陷类别（主要缺陷：检出框最多的类型）
            # 总数 = 检测的图片数
            file_results = project_info.get("file_results", [])
            computed_breakdown = {}
            for f_item in file_results:
                if not isinstance(f_item, dict):
                    continue
                img_grade = str(f_item.get("grade", "C"))
                # 无缺陷图片等级为 OK，跳过不参与 A/B/C 统计
                if img_grade == "OK":
                    continue
                if img_grade not in ("A", "B", "C"):
                    img_grade = "C"

                # 统计该图片上每种缺陷的检出框数量
                defect_counts = {}
                for det in f_item.get("detections", []):
                    if isinstance(det, dict):
                        cls_name = str(det.get("class_name", "")).strip()
                        if cls_name:
                            defect_counts[cls_name] = defect_counts.get(cls_name, 0) + 1
                # 也尝试从 defect_counts 字段获取
                for cls_name, cnt in f_item.get("defect_counts", {}).items():
                    if isinstance(cls_name, str) and cls_name.strip():
                        k = cls_name.strip()
                        if k not in defect_counts:
                            defect_counts[k] = int(cnt) if cnt else 0

                if not defect_counts:
                    continue

                # 选出检出框最多的缺陷类型作为该图片的主要缺陷
                primary = max(defect_counts, key=defect_counts.get)

                # 计入统计
                if primary not in computed_breakdown:
                    computed_breakdown[primary] = {"A": 0, "B": 0, "C": 0, "total": 0}
                computed_breakdown[primary][img_grade] = computed_breakdown[primary].get(img_grade, 0) + 1
                computed_breakdown[primary]["total"] = computed_breakdown[primary].get("total", 0) + 1

            # 优先使用前端传递的 defect_grade_breakdown，只有当前端未传递时才使用后端的计算结果
            if not defect_grade_breakdown and computed_breakdown:
                defect_grade_breakdown = computed_breakdown
            logger.info("缺陷等级分布(主要缺陷): %d 种缺陷, file_results=%d 条",
                        len(computed_breakdown), len(file_results))

            # ── 5c. 动态填充缺陷类别明细行 ──
            template_row_idx = None
            for ri, row in enumerate(t6.rows):
                # 将整行所有格子的内容拼接在一起，去掉空格
                row_text = "".join(c.text for c in row.cells).replace(" ", "").replace("　", "")
                # 如果这一行带有 {缺陷名称} 或者 {缺陷 （兼容全半角），才把它当做模板行
                # 但是我们要排除标题行，标题行通常不带大括号。
                if ("{缺陷名称}" in row_text or "｛缺陷名称｝" in row_text or "{缺陷" in row_text or "｛缺陷" in row_text):
                    template_row_idx = ri
                    break


                if template_row_idx is not None:
                    break

            if template_row_idx is not None and defect_grade_breakdown:
                template_row_el = t6.rows[template_row_idx]._tr
                defect_items = list(defect_grade_breakdown.items())

                template_tr_xml = copy.deepcopy(template_row_el)
                insert_after = template_row_el

                for idx, (cls_name, grade_info) in enumerate(defect_items):
                    if not isinstance(grade_info, dict):
                        grade_info = {"A": 0, "B": 0, "C": int(grade_info), "total": int(grade_info)}
                    cnt_a = grade_info.get("A", 0)
                    cnt_b = grade_info.get("B", 0)
                    cnt_c = grade_info.get("C", 0)
                    cnt_total = grade_info.get("total", cnt_a + cnt_b + cnt_c)
                    pct = f"{cnt_total / total_tested * 100:.1f}" if total_tested > 0 else "0"

                    if idx == 0:
                        target_tr = template_row_el
                    else:
                        new_tr = copy.deepcopy(template_tr_xml)
                        # 清除 w14:paraId 和 w14:textId（必须全文唯一）
                        # 必须清理 new_tr 自身 + 所有子元素
                        for key in list(new_tr.attrib.keys()):
                            if 'paraId' in key or 'textId' in key:
                                del new_tr.attrib[key]
                        for el in new_tr.xpath('.//*[@*]'):
                            for key in list(el.attrib.keys()):
                                if 'paraId' in key or 'textId' in key:
                                    del el.attrib[key]
                        insert_after.addnext(new_tr)
                        target_tr = new_tr
                        insert_after = new_tr

                    from docx.table import _Cell
                    from docx.oxml.ns import qn as _qn
                    tc_elements = target_tr.findall(_qn('w:tc'))
                    # 通过占位符关键字匹配填充（不依赖列索引）
                    for tc_el in tc_elements:
                        cell_obj = _Cell(tc_el, t6)
                        ct = cell_obj.text.replace(" ", "").replace("\u3000", "").replace("\n", "")
                        if "{缺陷名称}" in ct:
                            _set_cell_text(cell_obj, cls_name)
                        elif "{A类数量}" in ct:
                            _set_cell_text(cell_obj, str(cnt_a))
                        elif "{B类数量}" in ct:
                            _set_cell_text(cell_obj, str(cnt_b))
                        elif "{C类数量}" in ct:
                            _set_cell_text(cell_obj, str(cnt_c))
                        elif "{缺陷总数量}" in ct:
                            _set_cell_text(cell_obj, str(cnt_total))
                        elif "{占总测试比例}" in ct or "比例%" in ct:
                            _set_cell_text(cell_obj, pct)

                    logger.info("  缺陷行[%d]: %s A=%d B=%d C=%d total=%d pct=%s%%",
                                idx, cls_name, cnt_a, cnt_b, cnt_c, cnt_total, pct)

                logger.info("缺陷明细行: 填充了 %d 个缺陷类别", len(defect_items))

            elif template_row_idx is not None:
                # 无缺陷数据，清除模板行占位符
                for cell in t6.rows[template_row_idx].cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if "{" in r.text:
                                r.text = ""
                logger.info("无缺陷数据，已清除模板行占位符")
        else:
            logger.warning("未找到缺陷统计表")
    except Exception as exc:
        logger.warning("缺陷统计填充失败: %s", exc, exc_info=True)

    # ── 6. 填充 EL检测数据（图片表）──
    #
    # 策略：找到"附件1"或"附件 1"段落，在其后插入标题和表格。
    # 同时查找并删除模板中的原始 EL 数据占位表格。
    try:
        # 通过内容搜索定位 EL 数据表（包含"序号"+"组串号"或"EL照片"的表格）
        t7 = None
        for t in _orig_tables:
            if len(t.rows) > 0:
                header_text = "".join(c.text for c in t.rows[0].cells)
                if ("序号" in header_text and "组串" in header_text) or "EL照片" in header_text:
                    t7 = t
                    break
        # 查找占位符 {附件1 光伏组件电致发光（EL）检测数据} 段落作为插入点
        _appendix1_para = None
        _APPENDIX1_PH = "附件1 光伏组件电致发光"  # 占位符关键片段（避免全角半角不匹配）
        for p in doc.paragraphs:
            p_text = "".join(r.text for r in p.runs)
            if not p_text:
                p_text = p.text
            if _APPENDIX1_PH in p_text and "{" in p_text:
                _appendix1_para = p
                logger.info("找到附件1占位符段落: '%s'", p_text.strip()[:60])
                break
        if t7 is not None:
            file_results = project_info.get("file_results", [])
            # ── 按 image_filter 过滤 EL 图片 ──
            image_filter = project_info.get("image_filter", "all")
            if image_filter == "ng":
                file_results = [f for f in file_results if f.get("result") == "NG"]
                logger.info("EL图片按NG过滤: 过滤后%d条", len(file_results))
            elif image_filter == "ok":
                file_results = [f for f in file_results if f.get("result") == "OK"]
                logger.info("EL图片按OK过滤: 过滤后%d条", len(file_results))
            rotate_images = project_info.get("rotate_images", False)
            img_width_cm = float(project_info.get("img_width_cm", 10.0))
            img_height_cm = float(project_info.get("img_height_cm", 5.0))
            img_quality = int(project_info.get("img_quality", 85))
            img_quality = max(10, min(100, img_quality))
            annotation_style = _get_export_annotation_style(project_info)
            logger.info("EL检测数据: file_results数量=%d, rotate=%s, img_size=%.1fx%.1fcm, quality=%d, filter=%s",
                        len(file_results), rotate_images, img_width_cm, img_height_cm, img_quality, image_filter)

            # ── 6a. 分析目录层级 ──
            import os
            for f_item in file_results:
                img_path = f_item.get("path", "")
                dirname = os.path.dirname(img_path)
                folders = []
                current = dirname
                for _ in range(10):
                    parent = os.path.dirname(current)
                    if not current or current == parent:
                        break
                    folders.insert(0, os.path.basename(current))
                    current = parent
                if len(folders) >= 2:
                    f_item["_level1"] = folders[-2]
                    f_item["_level2"] = folders[-1]
                    f_item["_has_level1"] = True
                else:
                    f_item["_level1"] = ""
                    f_item["_level2"] = folders[-1] if folders else ""
                    f_item["_has_level1"] = False

            file_results.sort(key=lambda x: (x.get("_level1", ""), x.get("_level2", ""), x.get("name", "")))
            has_any_level1 = any(f.get("_has_level1", False) for f in file_results)
            # 调试日志：输出前3个文件的路径分析结果
            for fi, f_item in enumerate(file_results[:3]):
                logger.info("路径分析[%d]: path=%s, _level1=%s, _level2=%s, _has_level1=%s",
                            fi, f_item.get("path", "")[-60:], f_item.get("_level1"), f_item.get("_level2"), f_item.get("_has_level1"))

            if has_any_level1:
                all_results = list(file_results)
            else:
                ng_r = [f for f in file_results if f.get("result") == "NG"]
                ok_r = [f for f in file_results if f.get("result") == "OK"]
                ot_r = [f for f in file_results if f.get("result") not in ("NG", "OK")]
                all_results = ng_r + ok_r + ot_r

            # ── 6b. 按一级文件夹分组 ──
            from collections import OrderedDict
            groups = OrderedDict()
            for f_item in all_results:
                lv1 = f_item.get("_level1", "") if f_item.get("_has_level1") else ""
                groups.setdefault(lv1, []).append(f_item)

            # ── 6c. 保存 t7 的列宽，记录位置，删除原始 t7 ──
            t7_tbl_element = t7._tbl
            col_widths = []
            grid_cols = t7_tbl_element.findall(qn('w:tblGrid') + '/' + qn('w:gridCol'))
            if grid_cols:
                for gc in grid_cols:
                    w = gc.get(qn('w:w'))
                    col_widths.append(int(w) if w else 0)

            body = doc.element.body
            t7_prev = t7_tbl_element.getprevious()
            body.remove(t7_tbl_element)

            # ── 6d. 辅助：用 lxml 创建标题段落 ──
            from lxml import etree
            WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            def _make_heading_p(text: str) -> etree._Element:
                """创建一个 Heading 2 段落的 XML 元素，可在 Word 导航栏中显示。"""
                para = doc.add_paragraph(text, style='Heading 2')
                # 确保居中和加粗
                para.alignment = 1  # CENTER
                for run in para.runs:
                    run.font.bold = True
                p_el = para._element
                body.remove(p_el)
                return p_el

            def _make_page_break_p() -> etree._Element:
                """创建一个包含分页符的空段落。"""
                para = doc.add_page_break()
                p_el = para._element
                body.remove(p_el)
                return p_el

            # ── 6e. 辅助：创建 EL 数据表并填充 ──
            def _create_and_fill_table(items: list, start_seq: int = 1):
                """创建表格，填充数据，返回表格的 _tbl XML 元素。"""
                from docx.shared import Twips, Pt
                from docx.enum.table import WD_ALIGN_VERTICAL

                num_rows = 1 + len(items)
                logger.info("DEBUG: 创建表格行数=%d (items=%d + 1表头)", num_rows, len(items))
                tbl = doc.add_table(rows=num_rows, cols=5)
                tbl.style = 'Table Grid'
                tbl.autofit = False  # 禁用自动调整

                # 设置列宽（从模板 Table 7 复制）
                if col_widths and len(col_widths) == 5:
                    for i, w in enumerate(col_widths):
                        if w > 0:
                            tbl.columns[i].width = Twips(w)

                # 填充表头
                headers = ['序号', '组串号', '图片名称', '组件EL照片', '缺陷/数量']
                for ci, h in enumerate(headers):
                    cell = tbl.rows[0].cells[ci]
                    cell.text = h
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for p in cell.paragraphs:
                        p.alignment = 1  # CENTER
                        for run in p.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)

                # 填充数据行
                seq = start_seq
                total_items = len(items)
                for ri, f_item in enumerate(items):
                    if ri % 20 == 0:
                        logger.info("填充表格行: %d/%d", ri + 1, total_items)
                    row = tbl.rows[1 + ri]
                    # 序号列：居中
                    row.cells[0].text = str(seq)
                    for p in row.cells[0].paragraphs:
                        p.alignment = 1
                    # 组串号列：居中，10.5号字
                    row.cells[1].text = f_item.get("_level2", "")
                    for p in row.cells[1].paragraphs:
                        p.alignment = 1
                        for run in p.runs:
                            run.font.size = Pt(10.5)
                    # 图片名称列：居中，10.5号字
                    row.cells[2].text = f_item.get("name", "")
                    for p in row.cells[2].paragraphs:
                        p.alignment = 1
                        for run in p.runs:
                            run.font.size = Pt(10.5)
                    # EL照片列：图片，宽度适配列宽（约10cm）
                    _insert_el_image(row.cells[3], f_item, seq, rotate_images, Cm,
                                     img_width_cm, img_height_cm, img_quality, annotation_style)
                    # 缺陷/数量列：左对齐，9号字
                    _fill_defect_cell(row, f_item)
                    # 设置所有单元格垂直居中
                    for cell_obj in row.cells:
                        cell_obj.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    seq += 1

                # 从 body 末尾取出 tbl 元素（add_table 会追加 tbl + 空 p）
                tbl_element = tbl._tbl
                body.remove(tbl_element)
                # 不再删除尾随的空段落（Word文档严格要求必须以非表格元素结尾）
                return tbl_element

            # ── 6f. 在正确位置插入标题和表格 ──
            # 优先使用占位符段落作为插入点，并清除占位符文本
            if _appendix1_para is not None:
                insert_point = _appendix1_para._element
                # 清除占位符文本（保留段落元素作为锚点）
                for run in _appendix1_para.runs:
                    run.text = ""
                logger.info("使用占位符段落作为 EL 表格插入点")
            else:
                insert_point = t7_prev
                logger.info("使用原始 EL 数据表前一个元素作为插入点（未找到占位符）")
            logger.info("EL表格分组: has_any_level1=%s, groups=%d, group_keys=%s",
                         has_any_level1, len(groups), list(groups.keys()))
            _total_items_in_groups = sum(len(v) for v in groups.values())
            logger.info("DEBUG: groups中总图片数量=%d", _total_items_in_groups)
            if has_any_level1:
                # 有一级文件夹 → 每个分组一个标题 + 一个表格 + 分页符
                group_list = list(groups.items())
                for gi, (level1_name, group_items) in enumerate(group_list):
                    if level1_name:
                        heading_el = _make_heading_p(level1_name)
                        insert_point.addnext(heading_el)
                        insert_point = heading_el
                        logger.info("插入标题: '%s' (%d 张图片)", level1_name, len(group_items))
                    tbl_el = _create_and_fill_table(group_items, start_seq=1)
                    insert_point.addnext(tbl_el)
                    insert_point = tbl_el
                    # 每个分组表格后插入分页符（最后一组除外）
                    if gi < len(group_list) - 1:
                        pb_el = _make_page_break_p()
                        insert_point.addnext(pb_el)
                        insert_point = pb_el
            else:
                flat_items = []
                for items in groups.values():
                    flat_items.extend(items)
                tbl_el = _create_and_fill_table(flat_items, start_seq=1)
                insert_point.addnext(tbl_el)

    except Exception as exc:
        logger.warning("EL检测数据表填充失败: %s", exc, exc_info=True)

    # >>>>> CRITICAL GLOBAL PATCH START <<<<<
    try:
        if hasattr(doc, 'part') and hasattr(doc.part, 'package'):
            pic_count = 0
            fixed_count = 0
            for part in doc.part.package.parts:
                if part.content_type.endswith('xml') and hasattr(part, '_element') and part._element is not None:
                    for el in part._element.iter():
                        # Remove toxic tracking IDs
                        for k in list(el.attrib.keys()):
                            if 'paraId' in k or 'textId' in k:
                                del el.attrib[k]
                        
                        # Apply Sequential ID Override
                        if type(el.tag) is str:
                            if el.tag.endswith('}docPr'):
                                pic_count += 1
                                # Hard override docPr id and descr!
                                el.set('id', str(800000 + pic_count))
                                el.set('name', f'Picture {pic_count}')
                                if 'descr' not in el.attrib:
                                    el.set('descr', f'auto_img_{pic_count}')
                            
                            elif el.tag.endswith('}cNvPr'):
                                # Hard override cNvPr id!
                                el.set('id', str(900000 + pic_count))
                                fixed_count += 1

            # Write diagnostic file so we KNOW this ran
            debug_path = Path(__file__).parent.parent.parent / 'export_debug.txt'
            with open(debug_path, 'w', encoding='utf-8') as df:
                df.write(f'EXPORT PATCH RAN SUCCESSFULLY! Processed {pic_count} pictures. Fixed {fixed_count} zero-IDs.')
    except Exception as e:
        logger.warning(f"XML Clean Error: {e}", exc_info=True)
    # >>>>> CRITICAL GLOBAL PATCH END <<<<<

    # ── 7. 保存 ──
    target = Path(output_path).expanduser().resolve()
    target.parent.mkdir(parents=True, exist_ok=True)
    logger.info("开始保存文档到: %s", target)
    doc.save(str(target))
    file_size_mb = target.stat().st_size / (1024 * 1024)
    logger.info("模板导出成功: %s (%.1f MB)", target, file_size_mb)
    return {"message": "Word 报告导出成功（模板）", "output_path": target.as_posix()}


def _pv_get_module_value(module: dict, field_name: str) -> str:
    """Get pv module value, trying multiple key variants.
    
    Frontend keys use half-width parens with space: 'Pmax (Wp)', 'Voc (V)'
    Template/code field names use full-width parens: 'Pmax（Wp）', 'Voc（V）'
    This function tries all common variants to find a match.
    """
    # Build a list of candidate key names to try
    candidates = [field_name]
    # Variant 1: full-width → half-width (no space)
    v1 = field_name.replace("（", "(").replace("）", ")")
    candidates.append(v1)
    # Variant 2: full-width → half-width with space before '('
    v2 = field_name.replace("（", " (").replace("）", ")")
    candidates.append(v2)
    # Variant 3: half-width → full-width
    v3 = field_name.replace("(", "（").replace(")", "）")
    candidates.append(v3)
    # Variant 4: half-width with space → full-width (remove space before full-width paren)
    v4 = field_name.replace(" (", "（").replace(")", "）")
    candidates.append(v4)
    # Variant 5: half-width without space → half-width with space
    v5 = field_name.replace("(", " (")
    candidates.append(v5)
    
    for key in candidates:
        val = module.get(key, "")
        if val and str(val).strip():
            return str(val).strip()
    return ""


def _set_cell_text(cell, text: str):
    import re
    for p in cell.paragraphs:
        full = "".join(r.text for r in p.runs)
        if not full:
            continue
        replaced = re.sub(r'\{[^{}]*\}', text, full)
        if replaced != full:
            if not p.runs:
                p.add_run(replaced)
            else:
                p.runs[0].text = replaced
                for r in p.runs[1:]:
                    r.text = ""
    return


def _to_bool(value: Any, default: bool) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    if isinstance(value, str):
        lowered = value.strip().lower()
        if lowered in {"1", "true", "yes", "y", "on"}:
            return True
        if lowered in {"0", "false", "no", "n", "off"}:
            return False
    return default


def _to_int(value: Any, default: int) -> int:
    try:
        return int(float(value))
    except Exception:
        return default


def _to_float(value: Any, default: float) -> float:
    try:
        return float(value)
    except Exception:
        return default


def _get_export_annotation_style(project_info: dict[str, Any]) -> dict[str, Any]:
    stroke_width = _to_int(project_info.get("stroke_width"), 2)
    font_size = _to_int(project_info.get("font_size"), 16)
    return {
        "stroke_width": max(1, min(500, stroke_width)),  # 放宽上限，允许用户在界面填极大的数字
        "font_size": max(1, min(500, font_size)),      # 放宽上限和下限
        "show_boxes": _to_bool(project_info.get("show_boxes"), True),
        "show_labels": _to_bool(project_info.get("show_labels"), True),
        "show_confidence": _to_bool(project_info.get("show_confidence"), True),
    }


def _build_detection_items_from_file_result(f_item: dict[str, Any]) -> list[DetectionItem]:
    raw = f_item.get("detections", [])
    if not isinstance(raw, list):
        return []

    parsed: list[DetectionItem] = []
    for det in raw:
        if not isinstance(det, dict):
            continue

        box = det.get("box", {})
        if not isinstance(box, dict):
            continue

        x1 = _to_int(box.get("x1"), 0)
        y1 = _to_int(box.get("y1"), 0)
        x2 = _to_int(box.get("x2"), 0)
        y2 = _to_int(box.get("y2"), 0)
        x1, x2 = sorted((x1, x2))
        y1, y2 = sorted((y1, y2))
        if x2 <= x1 or y2 <= y1:
            continue

        class_id = _to_int(det.get("class_id"), 0)
        class_name = str(det.get("class_name", f"缺陷_{class_id}"))
        try:
            score = float(det.get("score", 0.0))
        except Exception:
            score = 0.0
        score = max(0.0, min(1.0, score))

        parsed.append(
            DetectionItem(
                class_id=class_id,
                class_name=class_name,
                score=score,
                x1=x1,
                y1=y1,
                x2=x2,
                y2=y2,
            )
        )

    return parsed


def _load_export_image_with_style(
    f_item: dict[str, Any],
    annotation_style: dict[str, Any],
    expected_output_long_side: int | None = None,
    preview_canvas_width: float | None = None,
    preview_canvas_height: float | None = None,
):
    """Load image for export and optionally redraw boxes with export style."""
    import cv2
    import numpy as np
    # 导出图中的标注需要比实时预览“适度放大”，以适应大图。
    # 用户在前端设置了“线段粗细”、“标签大小”等，我们需要按比例把它们映射到原图的分辨率。
    stroke_boost = 1.0
    font_boost = 1.0

    vis_path = str(f_item.get("visualization_path", "") or "")
    orig_path = str(f_item.get("path", "") or "")
    detections = _build_detection_items_from_file_result(f_item)

    use_redraw = bool(detections) and bool(orig_path) and Path(orig_path).exists()
    src_path = ""
    if use_redraw:
        src_path = orig_path
    elif vis_path and Path(vis_path).exists():
        src_path = vis_path
    elif orig_path and Path(orig_path).exists():
        src_path = orig_path

    if not src_path:
        return None, ""

    img_data = np.fromfile(src_path, dtype=np.uint8)
    if img_data.size == 0:
        return None, src_path
    image = cv2.imdecode(img_data, cv2.IMREAD_COLOR)
    if image is None:
        return None, src_path

    if use_redraw:
        h, w = image.shape[:2]
        src_long_side = max(w, h)
        # 用户在前端预览画布上设置的 stroke_width / font_size 是相对于预览画布尺寸的。
        # 导出时画到原图上，需要按照 "原图尺寸 / 预览画布尺寸" 的比例等比放大，
        # 这样导出的标注和用户在屏幕上预览时看到的视觉效果完全一致。
        preview_long = max(preview_canvas_width or 0, preview_canvas_height or 0)
        if preview_long > 0:
            style_scale = src_long_side / preview_long
        else:
            # 如果前端未传预览画布尺寸，用一个合理的默认值（假设预览约800px）
            style_scale = src_long_side / 800.0
        style_scale = max(1.0, style_scale)

        draw_stroke_width = max(
            1,
            int(round(annotation_style["stroke_width"] * style_scale * stroke_boost)),
        )
        draw_font_size = max(
            1,
            int(round(annotation_style["font_size"] * style_scale * font_boost)),
        )
        logger.info(
            "导出标注参数: 原图=%dx%d, 预览画布=%.0fx%.0f, style_scale=%.2f, "
            "用户stroke=%s, 用户font=%s -> draw_stroke=%d, draw_font=%d",
            w, h, preview_canvas_width or 0, preview_canvas_height or 0, style_scale,
            annotation_style["stroke_width"], annotation_style["font_size"],
            draw_stroke_width, draw_font_size,
        )

        rendered = image.copy()
        DefectDetectionEngine._draw_detections(
            rendered,
            detections,
            stroke_width=draw_stroke_width,
            font_size=draw_font_size,
            show_boxes=annotation_style["show_boxes"],
            show_labels=annotation_style["show_labels"],
            show_confidence=annotation_style["show_confidence"],
        )
        return rendered, src_path

    return image, src_path


def _insert_el_image(cell, f_item: dict, seq_num: int, rotate_images: bool, Cm,
                     img_width_cm: float = 10.0, img_height_cm: float = 5.0,
                     img_quality: int = 85,
                     annotation_style: dict[str, Any] | None = None):
    """向单元格插入 EL 图片。

    优先使用当前设置重绘后的标注图，并在写入 Word 前按目标显示尺寸缩放压缩，
    以控制文档体积并保持显示清晰。
    """
    if annotation_style is None:
        annotation_style = _get_export_annotation_style({})

    src_for_log = f_item.get("path", "")
    try:
        import cv2

        dpi = 150
        px_per_cm = dpi / 2.54
        target_w = int(img_width_cm * px_per_cm)
        target_h = int(img_height_cm * px_per_cm)
        target_long_side = max(target_w, target_h)

        image, src_path = _load_export_image_with_style(
            f_item,
            annotation_style,
            expected_output_long_side=target_long_side,
        )
        if image is None:
            cell.text = f_item.get("name", "")
            return
        src_for_log = src_path or src_for_log

        if rotate_images:
            image = cv2.rotate(image, cv2.ROTATE_90_COUNTERCLOCKWISE)

        # 150dpi 在清晰度和文档体积之间较平衡。
        h, w = image.shape[:2]
        if w > target_w or h > target_h:
            scale = min(target_w / w, target_h / h)
            new_w = int(w * scale)
            new_h = int(h * scale)
            image = cv2.resize(image, (new_w, new_h), interpolation=cv2.INTER_AREA)
            logger.debug("图片缩放: %dx%d -> %dx%d (seq=%d)", w, h, new_w, new_h, seq_num)

        # 统一输出为 JPEG，减小嵌入 Word 后的数据量。
        temp_name = f"el_export_{seq_num}.jpg"
        temp_path = Path(tempfile.gettempdir()) / temp_name
        _, buf = cv2.imencode('.jpg', image, [cv2.IMWRITE_JPEG_QUALITY, img_quality])
        buf.tofile(str(temp_path))

        # 清空单元格现有内容后插图。
        for p in cell.paragraphs:
            p.text = ""
            for r in p.runs:
                r.text = ""
        p0 = cell.paragraphs[0]
        run = p0.add_run()
        run.add_picture(str(temp_path), width=Cm(img_width_cm), height=Cm(img_height_cm))
        try:
            temp_path.unlink()
        except Exception:
            pass
    except Exception as img_exc:
        logger.warning("图片插入失败: %s - %s", src_for_log, img_exc)
        cell.text = f_item.get("name", "")

def _fill_defect_cell(row, f_item: dict):
    """填充数据行的缺陷/数量列（9号字体，左对齐）。"""
    from docx.shared import Pt
    if len(row.cells) <= 4:
        return
    cell5 = row.cells[4]
    defect_counts = f_item.get("defect_counts", {})
    if defect_counts and isinstance(defect_counts, dict) and any(v > 0 for v in defect_counts.values()):
        defect_lines = [f"{name}: {count}" for name, count in sorted(defect_counts.items()) if count > 0]
        # 清空并写入
        cell5.text = ""
        cell5.paragraphs[0].text = defect_lines[0] if defect_lines else ""
        for line in defect_lines[1:]:
            cell5.add_paragraph(line)
    elif f_item.get("result") == "OK":
        cell5.text = "无缺陷"
    else:
        cell5.text = f_item.get("result", "")
    # 设置所有段落为9号字体、左对齐
    for p in cell5.paragraphs:
        p.alignment = 0  # LEFT
        for run in p.runs:
            run.font.size = Pt(9)


def _fill_defect_row(table, row_idx: int, defect_name: str, defect_by_class: dict, total_tested: int):
    """填充缺陷统计行。"""
    if row_idx >= len(table.rows):
        return
    row = table.rows[row_idx]
    # 查找匹配的缺陷类别（模糊匹配）
    count = 0
    for cls_name, cnt in defect_by_class.items():
        if defect_name in cls_name or cls_name in defect_name:
            count += cnt
    # cells[4]: 数量, cells[5]: 判定类, cells[6]: 占比%
    if len(row.cells) > 4:
        _set_cell_text(row.cells[4], str(count))
    if len(row.cells) > 5:
        grade = "C" if count > 0 else "A"
        _set_cell_text(row.cells[5], grade)
    if len(row.cells) > 6:
        pct = f"{count / total_tested * 100:.1f}" if total_tested > 0 else "0"
        _set_cell_text(row.cells[6], pct)


def _export_word_fallback(request: dict[str, Any], output_path: str) -> dict[str, Any]:
    """程序化生成 Word 检测报告（严格按 NIT NNE RT-002 模板布局）。"""
    project_info = request.get("project_info", {})
    logger.info("Word报告导出开始: output=%s", output_path)
    logger.info("Word报告数据字段: %s", list(project_info.keys()))

    def _g(*keys: str, default: str = "") -> str:
        """按优先级查找字段值，支持多个中英文key。"""
        for k in keys:
            val = project_info.get(k)
            if val is not None and str(val).strip():
                return str(val).strip()
        return default

    try:
        from docx import Document
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        target = Path(output_path).expanduser().resolve()
        target.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        # 标题
        title = doc.add_heading("EL光伏组件缺陷检测报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 一、报告封面信息
        doc.add_heading("一、报告封面信息", level=1)
        info_fields = [
            ("报告编号", _g("报告编号", "report_code")),
            ("签发日期", _g("签发日期", "issue_date", "signing_date")),
            ("委托单位", _g("委托单位", "client_name")),
            ("检测单位", _g("检测单位", "test_unit", "testing_unit")),
            ("项目名称", _g("项目名称", "project_name")),
            ("项目地址", _g("项目地址", "project_address")),
        ]
        table = doc.add_table(rows=len(info_fields), cols=2, style="Table Grid")
        for i, (label, value) in enumerate(info_fields):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value

        # 二、技术尽调信息
        doc.add_heading("二、技术尽调信息", level=1)
        diligence_fields = [
            ("委托单位地址", _g("委托单位地址", "client_address")),
            ("尽调周期", _g("尽调周期", "diligence_period", "investigation_period")),
            ("样品来源", _g("样品来源", "sample_source")),
            ("抽样原则", _g("抽样原则", "sampling_principle")),
            ("参考标准", _g("参考标准", "reference_standard")),
            ("检测人员", _g("检测人员", "tester")),
            ("审核人员", _g("审核人员", "reviewer")),
            ("批准人员", _g("批准人员", "approver")),
            ("项目概述", _g("项目概述", "project_overview")),
        ]
        table_d = doc.add_table(rows=len(diligence_fields), cols=2, style="Table Grid")
        for i, (label, value) in enumerate(diligence_fields):
            table_d.rows[i].cells[0].text = label
            table_d.rows[i].cells[1].text = value

        # 插入图片（如果有）
        geo_img = _g("geographic_image_path")
        overall_img = _g("overall_image_path")
        if geo_img or overall_img:
            doc.add_heading("三、项目概述图片", level=1)
            if geo_img and Path(geo_img).exists():
                doc.add_paragraph("光伏厂区地理位置图：")
                try:
                    doc.add_picture(geo_img, width=Inches(5))
                except Exception as img_exc:
                    logger.warning("地理位置图插入失败(已跳过): %s — %s", geo_img, img_exc)
                    doc.add_paragraph(f"[图片无法插入: {Path(geo_img).name}]")
            if overall_img and Path(overall_img).exists():
                doc.add_paragraph("光伏厂区整体图：")
                try:
                    doc.add_picture(overall_img, width=Inches(5))
                except Exception as img_exc:
                    logger.warning("整体图插入失败(已跳过): %s — %s", overall_img, img_exc)
                    doc.add_paragraph(f"[图片无法插入: {Path(overall_img).name}]")

        # 四、电站基本信息
        doc.add_heading("四、电站基本信息", level=1)
        station_fields = [
            ("项目名称", _g("项目名称", "project_name")),
            ("电站地址", _g("电站地址", "station_address")),
            ("联系人", _g("联系人", "contact_person")),
            ("联系方式", _g("联系方式", "contact_phone")),
            ("业主单位名称", _g("业主单位名称", "owner_name")),
            ("设计单位名称", _g("设计单位名称", "design_unit")),
            ("EPC单位名称", _g("EPC单位名称", "epc_unit")),
            ("运维单位名称", _g("运维单位名称", "maintenance_unit")),
            ("电站直流安装容量", _g("电站直流安装容量", "dc_capacity")),
            ("分几期建设", _g("分几期建设")),
            ("建设状态", _g("建设状态")),
            ("电站类型", _g("电站类型")),
            ("土地类型", _g("土地类型")),
            ("土地现状", _g("土地现状")),
            ("占地面积", _g("占地面积")),
            ("设计倾角", _g("设计倾角")),
            ("建设成本", _g("建设成本")),
            ("地理坐标", _g("地理坐标")),
            ("地形地貌", _g("地形地貌")),
            ("水/土壤情况", _g("水/土壤情况")),
            ("电网接入方式", _g("电网接入方式")),
            ("并网电压等级", _g("并网电压等级")),
            ("并网接入距离", _g("并网接入距离")),
            ("主变容量", _g("主变容量")),
            ("是否限电", _g("是否限电")),
            ("限电调控方式", _g("限电调控方式")),
            ("并网时间", _g("并网时间")),
            ("上网电价", _g("上网电价")),
            ("桩基形式", _g("桩基形式")),
            ("支架形式", _g("支架形式")),
            ("单个组串组件数", _g("单个组串组件数")),
            ("组件固定方式", _g("组件固定方式")),
            ("组件下边缘距地高度", _g("组件下边缘距地高度")),
        ]
        # 过滤掉空值
        station_fields = [(k, v) for k, v in station_fields if v]
        if station_fields:
            table_s = doc.add_table(rows=len(station_fields), cols=2, style="Table Grid")
            for i, (label, value) in enumerate(station_fields):
                table_s.rows[i].cells[0].text = label
                table_s.rows[i].cells[1].text = value

        # 五、设备信息 — 光伏组件参数
        pv_modules = project_info.get("pv_modules", [])
        if pv_modules:
            doc.add_heading("五、设备信息 — 光伏组件参数", level=1)
            pv_field_pairs = [
                ("生产厂家", "型号"),
                ("类型", "Pmax (Wp)"),
                ("Voc (V)", "Vmp (V)"),
                ("Isc (A)", "Imp (A)"),
                ("组件尺寸 (mm)", "短路电流温度系数 (%/°C)"),
                ("功率温度系数 (%/°C)", "开路电压温度系数 (%/°C)"),
            ]
            for idx, mod in enumerate(pv_modules):
                if not isinstance(mod, dict):
                    continue
                title = f"光伏组件 {idx + 1}"
                # 标题行 + 6 数据行，每行 4 列 (label|value|label|value)
                tbl = doc.add_table(rows=len(pv_field_pairs) + 1, cols=4, style="Table Grid")
                # 合并标题行
                hdr = tbl.rows[0]
                hdr.cells[0].merge(hdr.cells[3])
                _set_cell_text(hdr.cells[0], title)
                for ri, (k1, k2) in enumerate(pv_field_pairs, 1):
                    row = tbl.rows[ri]
                    _set_cell_text(row.cells[0], k1)
                    _set_cell_text(row.cells[1], str(mod.get(k1, "")))
                    _set_cell_text(row.cells[2], k2)
                    _set_cell_text(row.cells[3], str(mod.get(k2, "")))
                if idx < len(pv_modules) - 1:
                    doc.add_paragraph()  # 间距

        # 六、检测结果统计
        doc.add_heading("六、检测结果统计", level=1)
        stats = [
            ("检测图片总数", str(project_info.get("total_images", 0))),
            ("NG图片数", str(project_info.get("ng_images", 0))),
            ("OK图片数", str(project_info.get("ok_images", 0))),
            ("缺陷类别数", str(project_info.get("defect_classes", 0))),
            ("缺陷总数", str(project_info.get("defect_total", 0))),
        ]
        table3 = doc.add_table(rows=len(stats), cols=2, style="Table Grid")
        for i, (label, value) in enumerate(stats):
            table3.rows[i].cells[0].text = label
            table3.rows[i].cells[1].text = value

        # 缺陷分类明细
        defect_by_class = project_info.get("defect_by_class", {})
        if defect_by_class:
            doc.add_heading("七、缺陷分类明细", level=1)
            table4 = doc.add_table(rows=len(defect_by_class) + 1, cols=3, style="Table Grid")
            table4.rows[0].cells[0].text = "缺陷类别"
            table4.rows[0].cells[1].text = "数量"
            table4.rows[0].cells[2].text = "占比"
            total_defects = max(1, sum(defect_by_class.values()))
            for i, (cls_name, count) in enumerate(defect_by_class.items(), 1):
                table4.rows[i].cells[0].text = cls_name
                table4.rows[i].cells[1].text = str(count)
                table4.rows[i].cells[2].text = f"{count / total_defects * 100:.1f}%"

        # 文件检测结果列表
        file_results = project_info.get("file_results", [])
        if file_results:
            doc.add_heading("八、检测文件明细", level=1)
            table5 = doc.add_table(rows=min(len(file_results), 500) + 1, cols=3, style="Table Grid")
            table5.rows[0].cells[0].text = "序号"
            table5.rows[0].cells[1].text = "文件名"
            table5.rows[0].cells[2].text = "检测结果"
            for i, f in enumerate(file_results[:500], 1):
                table5.rows[i].cells[0].text = str(i)
                table5.rows[i].cells[1].text = f.get("name", "")
                table5.rows[i].cells[2].text = f.get("result", "")

        doc.save(str(target))
        return {"message": "Word 报告导出成功", "output_path": target.as_posix()}
    except ImportError:
        raise HTTPException(status_code=400, detail="未安装 python-docx 库，请运行: pip install python-docx")
    except Exception as exc:
        logger.error("Word导出失败: %s", exc, exc_info=True)
        raise HTTPException(status_code=400, detail=f"Word导出失败: {type(exc).__name__}: {exc}") from exc


@router.post("/api/report/export_images")
def export_images(request: dict[str, Any]) -> dict[str, Any]:
    """导出带标注框的检测图片到指定目录（保持原始目录层级、原始格式、最高质量）。"""
    project_info = request.get("project_info", {})
    output_dir = request.get("output_dir", "")
    if not output_dir:
        raise HTTPException(status_code=400, detail="未指定输出目录")
    logger.info("图片导出开始: output_dir=%s", output_dir)
    try:
        import cv2
        import os

        target_dir = Path(output_dir).expanduser().resolve()
        target_dir.mkdir(parents=True, exist_ok=True)

        file_results = project_info.get("file_results", [])
        annotation_style = _get_export_annotation_style(project_info)
        preview_canvas_width = _to_float(project_info.get("preview_canvas_width"), 0.0)
        preview_canvas_height = _to_float(project_info.get("preview_canvas_height"), 0.0)

        # 计算所有原始路径的公共前缀，用于保持相对目录结构
        orig_paths = [f.get("path", "") for f in file_results if f.get("path")]
        common_prefix = ""
        if orig_paths:
            common_prefix = os.path.commonpath([os.path.dirname(p) for p in orig_paths if p])

        exported = 0
        for f_item in file_results:
            orig_path = f_item.get("path", "")
            name = f_item.get("name", "")
            if not name:
                continue
            try:
                image, _ = _load_export_image_with_style(
                    f_item,
                    annotation_style,
                    preview_canvas_width=preview_canvas_width,
                    preview_canvas_height=preview_canvas_height,
                )
                if image is None:
                    continue
                # 计算相对路径以保持目录层级
                if orig_path and common_prefix:
                    orig_dir = os.path.dirname(orig_path)
                    rel_dir = os.path.relpath(orig_dir, common_prefix)
                    if rel_dir == '.':
                        rel_dir = ''
                else:
                    rel_dir = ''
                out_dir = target_dir / rel_dir if rel_dir else target_dir
                out_dir.mkdir(parents=True, exist_ok=True)
                out_path = out_dir / name

                ext = Path(name).suffix.lower()
                if ext in ('.png',):
                    _, buf = cv2.imencode('.png', image, [cv2.IMWRITE_PNG_COMPRESSION, 0])
                elif ext in ('.tif', '.tiff'):
                    _, buf = cv2.imencode('.tiff', image)
                elif ext in ('.bmp',):
                    _, buf = cv2.imencode('.bmp', image)
                else:
                    _, buf = cv2.imencode('.jpg', image, [cv2.IMWRITE_JPEG_QUALITY, 100])
                buf.tofile(str(out_path))
                exported += 1
            except Exception as img_exc:
                logger.warning("图片导出失败: %s - %s", name, img_exc)
        return {"message": f"已导出 {exported} 张图片", "output_dir": target_dir.as_posix(), "exported_count": exported}
    except Exception as exc:
        logger.error("图片导出失败: %s", exc, exc_info=True)
        raise HTTPException(status_code=400, detail=f"图片导出失败: {exc}") from exc


@router.post("/api/report/export_excel")
def export_excel_report(request: dict[str, Any]) -> dict[str, Any]:
    """导出 Excel 检测报告（仅双Sheet结构，包含图片及名称）。"""
    project_info = request.get("project_info", {})
    output_path = request.get("output_path", "")
    if not output_path:
        raise HTTPException(status_code=400, detail="未指定输出路径")

    logger.info("Excel报告导出开始: output=%s", output_path)
    
    # 尺寸设定，后端接收到的如10.0cm转换为大致像素值
    img_width_cm = float(project_info.get("img_width_cm", 10.0))
    img_height_cm = float(project_info.get("img_height_cm", 5.0))
    # 1 cm 一般等于 37.8 像素（96DPI）
    px_width = int(img_width_cm * 37.8)
    px_height = int(img_height_cm * 37.8)
    
    # Openpyxl 的列宽大约为 Pixel / 7.4
    col_width_excel = px_width / 7.4
    # Openpyxl 的行高大约为 Pixel * 0.75
    row_height_excel = px_height * 0.75

    cols_per_row = int(project_info.get("img_cols", 10))
    quality = int(project_info.get("img_quality", 75))
    rotate_images = project_info.get("rotate_images", False)
    
    file_results = project_info.get("file_results", [])
    
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.drawing.image import Image as ExcelImage
        from openpyxl.utils import get_column_letter
        from PIL import Image as PILImage
        import io
        import math
        
        target = Path(output_path).expanduser().resolve()
        target.parent.mkdir(parents=True, exist_ok=True)

        wb = Workbook()
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin'),
        )
        header_fill = PatternFill(start_color="1D4ED8", end_color="1D4ED8", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=11)

        # --- 第一页：缺陷统计 ---
        ws1 = wb.active
        ws1.title = "缺陷统计"
        ws1.column_dimensions['A'].width = 8
        ws1.column_dimensions['B'].width = 20
        ws1.column_dimensions['C'].width = 12
        ws1.column_dimensions['D'].width = 12
        ws1.column_dimensions['E'].width = 30
        
        headers = ["序号", "缺陷类型", "数量", "百分比", "说明"]
        for col, h in enumerate(headers, 1):
            cell = ws1.cell(row=1, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center', vertical='center')
        
        defect_by_class = project_info.get("defect_by_class", {})
        total_defects = max(1, sum(defect_by_class.values()))
        
        for i, (cls_name, count) in enumerate(defect_by_class.items(), 2):
            ws1.cell(row=i, column=1, value=i - 1).border = thin_border
            ws1.cell(row=i, column=2, value=cls_name).border = thin_border
            ws1.cell(row=i, column=3, value=count).border = thin_border
            ws1.cell(row=i, column=4, value=f"{count / total_defects * 100:.1f}%").border = thin_border
            ws1.cell(row=i, column=5, value="").border = thin_border
            ws1.cell(row=i, column=1).alignment = Alignment(horizontal='center')
            ws1.cell(row=i, column=3).alignment = Alignment(horizontal='center')
            ws1.cell(row=i, column=4).alignment = Alignment(horizontal='center')
            
        # 合计行
        last_row = len(defect_by_class) + 2
        ws1.cell(row=last_row, column=1, value="合计").border = thin_border
        ws1.cell(row=last_row, column=2, value="").border = thin_border
        ws1.cell(row=last_row, column=3, value=sum(defect_by_class.values())).border = thin_border
        ws1.cell(row=last_row, column=4, value="100.0%").border = thin_border
        ws1.cell(row=last_row, column=5, value="").border = thin_border
        ws1.merge_cells(start_row=last_row, start_column=1, end_row=last_row, end_column=2)
        ws1.cell(row=last_row, column=1).alignment = Alignment(horizontal='center')
        ws1.cell(row=last_row, column=1).font = Font(bold=True)
        
        # --- 第二页：图片数据 ---
        ws2 = wb.create_sheet("图片数据")
        ws2.column_dimensions['A'].width = 15
        
        # 按照文件夹归类分离
        from collections import defaultdict
        groups = defaultdict(list)
        
        for f in file_results:
            path_str = f.get("path", "")
            if not path_str:
                continue
            path_obj = Path(path_str)
            parent_name = path_obj.parent.name if path_obj.parent.name else "未知目录"
            groups[parent_name].append(f)
            
        current_row = 1
        
        # 图片区域表头
        ws2.cell(row=current_row, column=1, value="目录名称").fill = header_fill
        ws2.cell(row=current_row, column=1).font = header_font
        ws2.cell(row=current_row, column=1).alignment = Alignment(horizontal='center', vertical='center')
        
        for col_idx in range(1, cols_per_row + 1):
            cell = ws2.cell(row=current_row, column=col_idx + 1, value=f"图片{col_idx}")
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            
        current_row += 1
        
        col_max_px = {}
        row_max_px = {}
        
        # 逐组排列
        for group_name, files in groups.items():
            needed_rows = math.ceil(len(files) / cols_per_row)
            # 奇数行为图，偶数行为字，每个逻辑行占用2个Excel行
            total_excel_rows_used = needed_rows * 2
            
            # 目录名纵向合并处理
            merge_start = current_row
            merge_end = current_row + total_excel_rows_used - 1
            if merge_end > merge_start:
                ws2.merge_cells(start_row=merge_start, start_column=1, end_row=merge_end, end_column=1)
                
            dir_cell = ws2.cell(row=merge_start, column=1, value=group_name)
            dir_cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            dir_cell.font = Font(bold=True)
            
            for row_idx in range(needed_rows):
                start_idx = row_idx * cols_per_row
                end_idx = min(start_idx + cols_per_row, len(files))
                row_files = files[start_idx:end_idx]
                img_row = current_row + row_idx * 2
                name_row = img_row + 1
                
                # Default height fallback
                ws2.row_dimensions[name_row].height = 20
                
                for col_offset, file_dict in enumerate(row_files):
                    col_index = col_offset + 2 # 从 B 列开始
                    img_path = file_dict.get("path")
                    if img_path and Path(img_path).exists():
                        try:
                            
                            import cv2
                            import numpy as np
                            
                            preview_w = float(project_info.get("preview_canvas_width", 800))
                            preview_h = float(project_info.get("preview_canvas_height", 600))
                            annotation_style = _get_export_annotation_style(project_info)
                            
                            drawn_img, actual_src = _load_export_image_with_style(
                                file_dict, annotation_style,
                                preview_canvas_width=preview_w,
                                preview_canvas_height=preview_h
                            )
                            
                            if drawn_img is not None:
                                if rotate_images:
                                    drawn_img = cv2.rotate(drawn_img, cv2.ROTATE_90_CLOCKWISE)
                                    
                                drawn_img_rgb = cv2.cvtColor(drawn_img, cv2.COLOR_BGR2RGB)
                                pil_img = PILImage.fromarray(drawn_img_rgb)
                                
                                orig_w, orig_h = pil_img.size
                                # Calculate display size bounded by px_width, px_height
                                ratio = min(px_width / orig_w, px_height / orig_h) if orig_w > 0 and orig_h > 0 else 1.0
                                disp_w = int(orig_w * ratio)
                                disp_h = int(orig_h * ratio)
                                
                                # Retain extremely high resolution, up to 6x the display size for clarity
                                render_ratio = min(1.0, (disp_w * 6.0) / orig_w) if orig_w > 0 else 1.0
                                render_w = int(orig_w * render_ratio)
                                render_h = int(orig_h * render_ratio)
                                
                                pil_img.thumbnail((render_w, render_h), PILImage.Resampling.LANCZOS)
                                
                                img_buffer = io.BytesIO()
                                pil_img.save(img_buffer, format='JPEG', quality=quality, optimize=True)
                                
                            xl_img = ExcelImage(img_buffer)
                            xl_img.width = disp_w
                            xl_img.height = disp_h
                            
                            cell_coord = ws2.cell(row=img_row, column=col_index).coordinate
                            ws2.add_image(xl_img, cell_coord)
                            
                            col_max_px[col_index] = max(col_max_px.get(col_index, 0), disp_w)
                            row_max_px[img_row] = max(row_max_px.get(img_row, 0), disp_h)
                            
                            # 正下方输入文件名
                            filename = Path(img_path).name
                            name_cell = ws2.cell(row=name_row, column=col_index, value=filename)
                            name_cell.alignment = Alignment(horizontal='center', vertical='center')
                            name_cell.font = Font(size=9)
                            
                        except Exception as e:
                            logger.error(f"处理图片 {img_path} 出错: {e}")
                            ws2.cell(row=img_row, column=col_index, value="[加载失败]")
                            
            current_row += total_excel_rows_used

        # 应用动态计算的列宽和行高，紧凑贴合图片
        for col_idx, max_px in col_max_px.items():
            # 调整除数使其完美贴合图片边缘 (7.65 在大多数 Windows 环境下最为精确)
            ws2.column_dimensions[get_column_letter(col_idx)].width = max_px / 7.65
        for row_idx, max_px in row_max_px.items():
            ws2.row_dimensions[row_idx].height = max_px * 0.75

        wb.save(str(target))
        return {"message": "Excel 报告导出成功", "output_path": target.as_posix()}
        
    except ImportError:
        raise HTTPException(status_code=400, detail="未安装 openpyxl 或 Pillow，请运行: pip install openpyxl Pillow")
    except Exception as exc:
        logger.error("Excel导出失败: %s", exc, exc_info=True)
        raise HTTPException(status_code=400, detail=f"Excel导出失败: {exc}") from exc


from fastapi import File, UploadFile
import io

@router.post("/api/utils/extract_gps")
async def extract_gps(file: UploadFile = File(...)):
    """Extract GPS metadata from an uploaded image using _ExifGpsHelper."""
    import tempfile
    import os
    try:
        content = await file.read()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        # Note: 'with' block already closes the file, saving it on Windows.
        # But if the file is locked or encoding is corrupted, cv2.imread fails.
        # Let's read the image via numpy instead!
        try:
            from PIL import Image
            import io
            with Image.open(io.BytesIO(content)) as img:
                w, h = img.size

            from app.exif_helper import _ExifGpsHelper
            helper = _ExifGpsHelper(tmp_path, w, h)
            
            if helper.has_gps:
                res = {
                    "success": True, 
                    "latitude": helper.cam_lat, 
                    "longitude": helper.cam_lon, 
                    "altitude": helper.cam_alt,
                    "gsd": getattr(helper, 'gsd', 0.0),
                    "focal_mm": getattr(helper, 'focal_mm', 0.0),
                    "img_w": w,
                    "img_h": h
                }
                logger.info("[extract_gps] SUCCESS: lat=%.6f, lon=%.6f, alt=%.1f, gsd=%.6f, focal=%.1f, img_w=%d, img_h=%d",
                            helper.cam_lat, helper.cam_lon, helper.cam_alt, 
                            getattr(helper, 'gsd', 0.0), getattr(helper, 'focal_mm', 0.0), w, h)
            else:
                res = {"success": False, "error": "No GPS data found"}
        finally:
            if os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass
                    
        return res
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}
